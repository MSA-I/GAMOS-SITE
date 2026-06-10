"""
Embed the 28 screenshots from D:\\GAMOS-screenshots-tmp\\ into a copy of the
user's Word document, anchored to the matching section heading.

Source doc:  C:\\Users\\art1\\Desktop\\GAMOS-SITE-תוכן.docx  (READ-ONLY)
Target doc:  C:\\Users\\art1\\Desktop\\GAMOS-SITE-תוכן + צילומי מסך.docx

For each entry in HEADING_TO_IMAGE, we find the paragraph whose text matches
exactly, then insert a new paragraph immediately after with the image.
"""

import os, sys, shutil
from copy import deepcopy
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = r"D:\GAMOS-screenshots-tmp"
SRC_DOC  = r"C:\Users\art1\Desktop\GAMOS-SITE-תוכן.docx"
DEST_DOC = r"C:\Users\art1\Desktop\GAMOS-SITE-תוכן + צילומי מסך.docx"

# Map heading text (verbatim from doc) → screenshot filename in SHOTS dir.
HEADING_TO_IMAGE = {
    # === עמוד הבית (Heading 2) ===
    "Hero — פתיחה  ·  #hero":                              "01-hero.png",
    "לאונג' — קרוסלה תלת-ממדית  ·  #lounge":               "02-lounge.png",
    "קולינריה — סקרול scene + 9 מנות  ·  #culinary":      "03-culinary.png",
    "שבתות חתן — פרלקס אופקי  ·  #shabbat-chatan":         "04-shabbat-chatan.png",
    "חדרי נופש — 10 חדרים  ·  #rooms":                     "05-rooms.png",
    "Marquee — סרט נע (פס מעבר)  ·  marquee":              "06-marquee.png",
    "אודות  ·  #about":                                    "07-about.png",
    "המלצות — סליידר  ·  #testimonials":                  "08-testimonials.png",
    "גלריה — מוזאיקה 8 תמונות  ·  #gallery":              "09-gallery.png",
    "סוגי אירועים — אקורדיון  ·  #events":                 "10-events.png",
    "כשרות  ·  #kosher":                                   "11-kosher.png",
    "צור קשר — טופס + פרטי קשר  ·  #contact":              "12-contact.png",
    "מסלולי הגעה — 3 מפות  ·  #routes":                    "13-routes.png",

    # === עמודי משפט (Heading 2) ===
    "מדיניות פרטיות · /legal/privacy.html":                "15-privacy.png",
    "תנאי שימוש · /legal/terms.html":                      "16-terms.png",
    "הצהרת נגישות · /legal/accessibility.html":            "17-accessibility.png",

    # === אולם האירועים — קלפים (Heading 3) ===
    "קלף 01 · היכל האואזיס · The Oasis Hall":             "18-oasis-01.png",
    "קלף 02 · סיפון המיראז' · The Mirage Deck":           "19-oasis-02.png",
    "קלף 03 · במת המים · The Water Altar":                "20-oasis-03.png",
    "קלף 04 · קבלת פנים שקיעה · Sunset Reception":        "21-oasis-04.png",
    "קלף 05 · משתה המדבר · Desert Banquet":               "22-oasis-05.png",
    "קלף 06 · פרטי האומנות · Curated Details":            "23-oasis-06.png",

    # === ריזורט — קלפים (Heading 3) ===
    "קלף 01 · תעלת האורות · The Dusk Canal":              "24-lumina-01.png",
    "קלף 02 · התרחבות הספירות · Floating Harmony":         "25-lumina-02.png",
    "קלף 03 · שער הדמדומים · The Dusk Portal":            "26-lumina-03.png",
    "קלף 04 · אופקים זוהרים · Glowing Horizons":          "27-lumina-04.png",
    "קלף 05 · הגסטרונומיה הזוהרת · Lumina Feast":         "28-lumina-05.png",
    "קלף 06 · חלקיקי הזוהר · Celestial Textures":         "29-lumina-06.png",
}


def insert_image_after(paragraph, image_path):
    """Insert a new bidi-RTL center-aligned paragraph with image right after `paragraph`."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)

    # Programmatically import the new paragraph back into python-docx
    # by binding it to a Paragraph object
    from docx.text.paragraph import Paragraph
    p_obj = Paragraph(new_p, paragraph._parent)
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mark it as RTL bidi (so later inserts stay aligned with the rest of the doc)
    pPr = p_obj._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)

    run = p_obj.add_run()
    # 6.0" wide ≈ 15.24cm, fits within standard A4/Letter margins
    run.add_picture(image_path, width=Inches(6.0))
    return p_obj


def is_heading(p, levels=("Heading 2", "Heading 3", "Heading 4")):
    return p.style and p.style.name in levels


def main():
    if not os.path.exists(SRC_DOC):
        print(f"Source doc not found: {SRC_DOC}", file=sys.stderr); sys.exit(1)

    # Validate every image referenced exists
    missing = []
    for h, img in HEADING_TO_IMAGE.items():
        path = os.path.join(SHOTS, img)
        if not os.path.exists(path):
            missing.append(img)
    if missing:
        print(f"MISSING screenshots: {missing}", file=sys.stderr); sys.exit(1)

    # Open the source, save as the destination, then re-open the dest for editing
    shutil.copy2(SRC_DOC, DEST_DOC)
    print(f"Copied: {SRC_DOC} → {DEST_DOC}")

    doc = Document(DEST_DOC)

    inserted = 0
    expected = len(HEADING_TO_IMAGE)
    used = set()

    # Walk paragraphs in document order. For each heading whose text is in
    # the map, insert image right after.
    for p in list(doc.paragraphs):
        if not is_heading(p):
            continue
        text = p.text.strip()
        if text in HEADING_TO_IMAGE and text not in used:
            img_path = os.path.join(SHOTS, HEADING_TO_IMAGE[text])
            insert_image_after(p, img_path)
            used.add(text)
            inserted += 1
            print(f"  [{inserted:2d}/{expected}] {text[:60]}  →  {HEADING_TO_IMAGE[text]}")

    not_found = set(HEADING_TO_IMAGE) - used
    if not_found:
        print("\nWARNING — these headings were NOT found in the doc:", file=sys.stderr)
        for h in not_found:
            print(f"  · {h!r}", file=sys.stderr)

    doc.save(DEST_DOC)
    print(f"\nSaved: {DEST_DOC}")
    print(f"Inserted {inserted}/{expected} images.")


if __name__ == "__main__":
    main()
