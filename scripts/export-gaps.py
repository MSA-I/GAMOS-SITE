# -*- coding: utf-8 -*-
"""
export-gaps.py — בונה מסמך Word ובו *רק* רשימת התוכן החסר ממרון ירקוני,
היכן כל טקסט יושב באתר, ומה בדיוק צריך לספק — עם צילום המסך של כל סקציה
מוטמע מתוך הקובץ הקיים "GAMOS-SITE-תוכן + צילומי מסך.docx".

Usage:
    python scripts/export-gaps.py
"""

import io
import zipfile
from pathlib import Path

from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Brand palette (Constitution §5) ─────────────────────────────────────────
BRASS       = "CFAE83"
BRASS_DEEP  = "8B6F46"
COCOA       = "534133"
IVORY       = "F5EFE6"
INK_DEEP    = "1A1410"
ACCENT_ROSE = "B8576F"
MIST        = "E8DFD3"

EXPORT_DIR = Path(r"D:\משה פרוייקטים\GAMOS-DOCS\מסמכים מרון ירקוני\אקספורט")
SRC_DOCX   = EXPORT_DIR / "GAMOS-SITE-תוכן + צילומי מסך.docx"
OUT_DOCX   = EXPORT_DIR / "גאמוס — מה חסר להשלמת הטקסטים.docx"


# ── Pull screenshots out of the existing docx ───────────────────────────────
def load_screenshots():
    """Returns {media_name: PNG bytes} for the media we need."""
    shots = {}
    with zipfile.ZipFile(SRC_DOCX) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and name.endswith(".png"):
                shots[Path(name).name] = z.read(name)
    return shots


SHOTS = load_screenshots()


# ── RTL helpers (mirrors scripts/export-content.py) ─────────────────────────
def _set_para_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)


def _set_run_rtl(run):
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _add_bottom_border(paragraph, color_hex, size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h(doc, text, level, color_hex=None, size_pt=None, bottom_border=False):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"] if level <= 9 else doc.styles["Heading 9"]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_rtl(p)
    run = p.add_run(text)
    _set_run_rtl(run)
    if size_pt:
        run.font.size = Pt(size_pt)
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    run.font.bold = True
    if bottom_border:
        _add_bottom_border(p, BRASS)
    return p


def add_para(doc, text, color_hex=None, size_pt=11, italic=False, bold=False,
             indent_in=0, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_rtl(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.35
    if indent_in:
        p.paragraph_format.right_indent = Inches(indent_in)
    run = p.add_run(text)
    _set_run_rtl(run)
    run.font.size = Pt(size_pt)
    run.font.italic = italic
    run.font.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return p


def add_eyebrow(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_rtl(p)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(text.upper())
    _set_run_rtl(run)
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BRASS_DEEP)


def add_bullets(doc, items, color_hex=INK_DEEP, size_pt=11):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_rtl(p)
        run = p.add_run(item)
        _set_run_rtl(run)
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor.from_string(color_hex)


def add_callout(doc, label, text, fill_hex, label_color):
    """A single-cell shaded box used for 'what to provide' / conflicts."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    _set_cell_shading(cell, fill_hex)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_rtl(cp)
    r1 = cp.add_run(label + "  ")
    _set_run_rtl(r1)
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor.from_string(label_color)
    r2 = cp.add_run(text)
    _set_run_rtl(r2)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(INK_DEEP)
    add_para(doc, "", space_after=4)


def add_data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        if col_widths:
            cell.width = Inches(col_widths[j])
        _set_cell_shading(cell, BRASS)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_rtl(cp)
        run = cp.add_run(h)
        _set_run_rtl(run)
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(INK_DEEP)
    for i, row in enumerate(rows, start=1):
        zebra = i % 2 == 0
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            if col_widths:
                cell.width = Inches(col_widths[j])
            _set_cell_shading(cell, MIST if zebra else "FFFFFF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_para_rtl(cp)
            run = cp.add_run(str(val))
            _set_run_rtl(run)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(INK_DEEP)
    add_para(doc, "", space_after=6)


def add_screenshot(doc, media_name, width_in=6.0):
    """Embeds a screenshot from the source docx, centered, with a thin caption."""
    if media_name not in SHOTS:
        add_para(doc, f"[צילום מסך חסר: {media_name}]", color_hex=ACCENT_ROSE,
                 size_pt=9, italic=True)
        return
    stream = io.BytesIO(SHOTS[media_name])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(stream, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_rtl(cap)
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.add_run("צילום מסך מהאתר החי — להמחשת מיקום הטקסט")
    _set_run_rtl(cr)
    cr.font.size = Pt(8)
    cr.font.italic = True
    cr.font.color.rgb = RGBColor.from_string(COCOA)


# ── Document setup ──────────────────────────────────────────────────────────
def configure_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "David"
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:cs"), "David")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:ascii"), "Calibri")


def configure_section_rtl(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def add_footer(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_rtl(p)
        run = p.add_run("© 2026 גאמוס · gamos.co.il · רשימת תוכן חסר להשלמה")
        _set_run_rtl(run)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COCOA)


# ── One reusable block per missing section ──────────────────────────────────
def gap_block(doc, *, num, name, dom, eyebrow, what_needed, slots_headers,
              slots_rows, slots_widths, screenshot, note=None):
    add_eyebrow(doc, eyebrow)
    add_h(doc, f"{num}. {name}", level=2, color_hex=COCOA, size_pt=16,
          bottom_border=True)
    add_para(doc, f"מיקום בקוד: {dom}", color_hex=BRASS_DEEP, size_pt=10,
             italic=True, space_after=6)
    add_callout(doc, "מה צריך לספק:", what_needed, MIST, ACCENT_ROSE)
    if note:
        add_callout(doc, "שימו לב:", note, "F3E7D9", BRASS_DEEP)
    add_screenshot(doc, screenshot)
    add_para(doc, "הקופי הקיים כיום (להחלפה / כמבנה למילוי):",
             color_hex=COCOA, size_pt=10, bold=True, space_after=4)
    add_data_table(doc, slots_headers, slots_rows, col_widths=slots_widths)
    doc.add_page_break()


def build():
    doc = Document()
    configure_default_style(doc)
    configure_section_rtl(doc)
    add_footer(doc)

    # ── Cover ────────────────────────────────────────────────────────────
    add_h(doc, "גאמוס — מה חסר להשלמת עדכון הטקסטים", level=1,
          color_hex=BRASS_DEEP, size_pt=26, bottom_border=True)
    add_para(doc, "רשימת התוכן שעדיין צריך לקבל ממרון ירקוני, היכן כל טקסט "
                  "יושב באתר, ומה בדיוק צריך לספק.",
             color_hex=COCOA, size_pt=13, space_after=2)
    add_para(doc, "כל סקציה מלווה בצילום מסך מהאתר החי להמחשת המיקום.",
             color_hex=COCOA, size_pt=13, space_after=2)
    add_para(doc, "gamos.co.il  ·  די זהב 7, פארק ישראל, מעלה אדומים",
             color_hex=BRASS_DEEP, size_pt=11, italic=True, space_after=18)

    add_para(doc, "סיכום מהיר:", color_hex=INK_DEEP, size_pt=12, bold=True,
             space_after=2)
    add_bullets(doc, [
        "6 בלוקי תוכן מרכזיים חסרים: לאונג', שבתות חתן, חדרי נופש, המלצות, "
        "מנות קולינריה, ו-12 קלפי האולם/ריזורט.",
        "2 בלוקים משניים: תיאורי גלריה (alt) ומסלולי הגעה.",
        "4 סתירות שצריך להכריע (קיבולת, מספר חדרים, סטטיסטיקות, שדה הטופס).",
    ], color_hex=INK_DEEP, size_pt=11)
    doc.add_page_break()

    # ── PART A — high-priority content ───────────────────────────────────
    add_h(doc, "חלק א׳ — תוכן מרכזי חסר", level=1, color_hex=BRASS_DEEP,
          size_pt=22, bottom_border=True)
    add_para(doc, "", space_after=4)

    gap_block(
        doc, num="1", name="לאונג׳", dom="index.html · סקציה #lounge",
        eyebrow="לאונג'",
        what_needed="כותרת ראשית + פתיח קצר, ולכל אחד מ-8 פריימי הקרוסלה: "
                    "שם קצר (2-3 מילים) + משפט תיאור אחד.",
        slots_headers=["#", "שם הפריים (כיום)", "תיאור (כיום)"],
        slots_rows=[
            ("1", "קבלת הפנים", "כוסית ראשונה, השיחות מתחממות."),
            ("2", "שעת הזהב", "חצי שעה לפני שקיעה, האור הכי טוב לתמונות."),
            ("3", "שקיעה", "השמש יורדת מעל המדבר, הבר מתחיל לעבוד."),
            ("4", "אווירת לילה", "האור יורד, השיחות מעמיקות, הצללים חמים."),
            ("5", "בעיצומו של הערב", "האולם מלא, הצחוק חוזר מהקירות."),
            ("6", "אחרי חצות", "החתן והכלה על הרחבה, מי שלא רוקד יושב כאן."),
            ("7", "בחוץ", "פינת ישיבה תחת כיפת השמיים, שקט באמצע הערב."),
            ("8", "ערב סתיו", "קריר בחוץ, חמים בפנים. התקופה האהובה עלינו."),
        ],
        slots_widths=[0.4, 1.8, 4.2],
        screenshot="image2.png",
        note="כיום הכותרת: \"קבלת פנים שלא רוצים שתסתיים\" · הפתיח: "
             "\"הבר של גאמוס. כל שעה בערב נראית קצת אחרת.\"",
    )

    gap_block(
        doc, num="2", name="שבתות חתן", dom="index.html · סקציה #shabbat-chatan",
        eyebrow="העונה",
        what_needed="5 כותרות קצרות (משפט אחד כל אחת) שמתגלות בזו אחר זו תוך "
                    "כדי גלילת המסך. אין תיאורים ואין כפתורים — רק 5 כותרות.",
        slots_headers=["פאנל", "כותרת (כיום)"],
        slots_rows=[
            ("1", "צאו מהקצב היומיומי"),
            ("2", "שולחנות ערוכים בכבוד"),
            ("3", "שירה שעולה מן הלב"),
            ("4", "סעודות שלוש על נפלאות העונה"),
            ("5", "שבת של נשמה, ביחד"),
        ],
        slots_widths=[0.8, 5.6],
        screenshot="image4.png",
    )

    gap_block(
        doc, num="3", name="חדרי נופש", dom="index.html · סקציה #rooms",
        eyebrow="חדרי אירוח",
        what_needed="לכל חדר: שם + משפט תיאור אחד. (ראו סתירת מספר החדרים "
                    "בחלק ג׳ — 10 באתר מול 33 במסמך האודות.)",
        slots_headers=["מס'", "שם החדר (כיום)", "תיאור (כיום)"],
        slots_rows=[
            ("01", "סוויטת הגן", "נפתחת ישר אל הגן, חלון פנורמי, אמבט עומד."),
            ("02", "חדר הזוג", "קטן, רך, שקט. בדיוק מה שצריך אחרי האירוע."),
            ("03", "סוויטת הגג", "מהקומה השלישית רואים הכל. שקט של גובה."),
            ("04", "חדר דה-לוקס", "מרפסת פרטית אל הירוק, פינת קפה בבוקר."),
            ("05", "חדר קלאסי", "פשוט, נקי, מסודר. לישון, להתעורר, להמשיך."),
            ("06", "סוויטת הספא", "ג'קוזי בחדר עצמו. תוכניות לבוקר משתבשות."),
            ("07", "חדר משפחה", "שתי מיטות, מקום לילדים, אף אחד לא נדחק."),
            ("08", "סוויטת הכלה", "החדר הכי גדול, הכי שקט, הכי מהקצה."),
            ("09", "חדר פרמיום", "קצת יותר מרווח, מקלחת גשם, חלוקים."),
            ("10", "סוויטת בוטיק", "החביב עלינו. אינטימי, פרטי, נסתר."),
        ],
        slots_widths=[0.5, 1.7, 4.2],
        screenshot="image5.png",
        note="כיום הכותרת: \"חדרים לשינה רגועה בסוף ערב מושלם\".",
    )

    gap_block(
        doc, num="4", name="המלצות לקוחות", dom="index.html · סקציה #testimonials",
        eyebrow="המלצות",
        what_needed="8 ציטוטי לקוחות אמיתיים (עם אישור פרסום), ולכל ציטוט: "
                    "שם החותם + סוג האירוע. הציטוטים הנוכחיים הם דמה.",
        slots_headers=["#", "חתימה (כיום)", "סוג אירוע (כיום)"],
        slots_rows=[
            ("1", "נועה ויונתן", "חתונה"),
            ("2", "משפחת לוי", "בר־מצווה"),
            ("3", "שירה, מנהלת רווחה", "אירוע חברה"),
            ("4", "תמר ויעקב", "בת־מצווה"),
            ("5", "אבי, אורח באירוע", "אירוע פרטי"),
            ("6", "מיכל ודניאל", "חתונת סופ\"ש"),
            ("7", "רחל, אם הכלה", "חתונה כפולה"),
            ("8", "עדי וערן", "ברית מילה"),
        ],
        slots_widths=[0.4, 2.6, 3.4],
        screenshot="image8.png",
        note="כיום הכותרת: \"מה מספרים על האירוע אצלנו\". כל שורה למעלה "
             "מייצגת ציטוט אחד שצריך טקסט מלא + חתימה + סוג.",
    )

    gap_block(
        doc, num="5", name="מנות קולינריה", dom="index.html · סקציה #culinary",
        eyebrow="המטבח שלנו",
        what_needed="לכל אחת מ-9 המנות בגלריה: תגית קצרה + שם מנה + משפט תיאור. "
                    "(הנרטיב הכללי של הקולינריה כבר התקבל ממך.)",
        slots_headers=["#", "תגית (כיום)", "שם המנה (כיום)", "תיאור (כיום)"],
        slots_rows=[
            ("1", "פתיחה", "מנת פתיחה", "פתיח עדין שמכין את החיך."),
            ("2", "ראשונה", "מנה ראשונה", "מרכיבים טריים בהגשה מדויקת."),
            ("3", "עיקרית", "מנה עיקרית", "לב הארוחה: עשיר, מאוזן."),
            ("4", "שף", "צלחת השף", "יצירה אישית של הצוות."),
            ("5", "קינוח", "קינוח", "סיום מתוק ומעודן."),
            ("6", "מתוקים", "פינת המתוקים", "שפע צבעוני שחותם את החגיגה."),
            ("7", "חדש", "מנה עונתית", "מהמטבח של השבוע. (טקסט זמני)"),
            ("8", "חדש", "מנת בית", "קלאסיקה של גאמוס. (טקסט זמני)"),
            ("9", "חדש", "סיום ערב", "המנה האחרונה לפני סוף הערב. (זמני)"),
        ],
        slots_widths=[0.3, 1.0, 1.5, 3.6],
        screenshot="image3.png",
        note="3 המנות האחרונות (7-9) נושאות טקסט זמני שממתין לקופי שלך.",
    )

    gap_block(
        doc, num="6", name="קלפי האולם והריזורט (12 קלפים)",
        dom="halls/src/projectsData.ts · דפי /halls/dist/oasis/ + /halls/dist/lumina/",
        eyebrow="גלריית האולמות",
        what_needed="כותרת + תת-כותרת לכל אחד מ-12 הקלפים (6 אולם + 6 ריזורט). "
                    "מבנה הקלף ופרטי הצילום הקיימים נשארים — צריך רק את הכותרות.",
        slots_headers=["מתחם", "מס'", "כותרת (כיום)", "תת-כותרת (כיום)"],
        slots_rows=[
            ("אולם", "01", "היכל האואזיס", "סעודת גורמה תחת כיפת השמיים"),
            ("אולם", "02", "סיפון המיראז'", "בריכת אינסוף מול הדיונות"),
            ("אולם", "03", "במת המים", "חופה ארכיטקטונית צפה על המים"),
            ("אולם", "04", "קבלת פנים שקיעה", "מתחם טקסים פנורמי"),
            ("אולם", "05", "משתה המדבר", "סעודת אבירים בעיצוב ארצי"),
            ("אולם", "06", "פרטי האומנות", "ניואנסים של עושר וטבע"),
            ("ריזורט", "01", "תעלת האורות", "המדבר וחוג הספירות הזוהרות"),
            ("ריזורט", "02", "התרחבות הספירות", "ספירות צפות ומפזרות זוהר"),
            ("ריזורט", "03", "שער הדמדומים", "טבילה של מים והשתקפויות שמיים"),
            ("ריזורט", "04", "אופקים זוהרים", "הרי המדבר מאחורי כדורי הזוהר"),
            ("ריזורט", "05", "הגסטרונומיה הזוהרת", "סעודה מוקפת בספירות חמימות"),
            ("ריזורט", "06", "חלקיקי הזוהר", "פרטי העיצוב של ההיכל השני"),
        ],
        slots_widths=[0.9, 0.4, 2.0, 3.1],
        screenshot="image14.png",
        note="הקלפים נבנים מחדש דרך build:halls אחרי העדכון. למטה צילום "
             "של קלף אולם לדוגמה; קלפי הריזורט בנויים זהה.",
    )

    # ── PART B — secondary content ───────────────────────────────────────
    add_h(doc, "חלק ב׳ — תוכן משני חסר", level=1, color_hex=BRASS_DEEP,
          size_pt=22, bottom_border=True)
    add_para(doc, "", space_after=4)

    gap_block(
        doc, num="7", name="תיאורי גלריה (Alt לתמונות)",
        dom="index.html · סקציה #gallery",
        eyebrow="גלריה",
        what_needed="8 תיאורי תמונה קצרים (alt) — לנגישות וקוראי-מסך. "
                    "אופציונלי: כיתוב גלוי לכל תמונה.",
        slots_headers=["#", "Alt (כיום)"],
        slots_rows=[
            ("1", "מבט רחב על אולם האירועים המפואר"),
            ("2", "אורחים חוגגים ברגע של שמחה"),
            ("3", "תאורה חמה על במת האירוע"),
            ("4", "מנה מעוטרת בעדינות על השולחן"),
            ("5", "עיצוב שולחן אלגנטי בפרטי פרטים"),
            ("6", "רחבת ריקודים גועשת בערב"),
            ("7", "רגע פורמלי מטקס האירוע"),
            ("8", "פרט אישי ואינטימי מהחגיגה"),
        ],
        slots_widths=[0.4, 6.0],
        screenshot="image9.png",
    )

    gap_block(
        doc, num="8", name="מסלולי הגעה", dom="index.html · סקציה #routes",
        eyebrow="איך מגיעים",
        what_needed="לכל אחת מ-3 המפות: תיאור מסלול קצר (מאיזה כיוון ודרך אילו "
                    "יישובים). אופציונלי לעדכן את רשימת היישובים.",
        slots_headers=["כיוון", "תיאור המסלול (כיום)"],
        slots_rows=[
            ("מירושלים", "תלפיות, הר חומה, מרכז העיר"),
            ("מהמרכז", "רמות, גבעת זאב, גבעון"),
            ("מצפון", "בית אל, עפרה, רימונים, אדם"),
        ],
        slots_widths=[1.2, 5.2],
        screenshot="image13.png",
    )

    # ── PART C — conflicts to resolve ────────────────────────────────────
    add_h(doc, "חלק ג׳ — סתירות להכרעה", level=1, color_hex=ACCENT_ROSE,
          size_pt=22, bottom_border=True)
    add_para(doc, "ארבע נקודות שבהן הקופי החדש סותר את המידע שכבר באתר. "
                  "צריך החלטה אחת ברורה לכל אחת:",
             color_hex=INK_DEEP, size_pt=11, space_after=8)

    add_data_table(
        doc,
        ["נושא", "באתר כיום", "במסמכי מרון", "צריך להכריע"],
        [
            ("קיבולת אורחים", "1,000+", "1,000 (מתחמים) / 500 (אירועים עסקיים)",
             "מספר אחד, או הבחנה לפי סוג אירוע"),
            ("מספר חדרים", "10 חדרים", "33 סוויטות",
             "המספר הנכון + שמות/תיאורים"),
            ("סטטיסטיקות אודות", "8 שנים · 200+ אירועים/שנה", "\"עשרות שנים\" ניסיון",
             "מספרים מאומתים"),
            ("טופס יצירת קשר", "כפתור \"שלחו דרך WhatsApp\", ללא שדה חברה",
             "כפתור \"שליחת פרטים\" + שדה \"שם החברה\"",
             "לאשר נוסח כפתור + הוספת שדה"),
        ],
        col_widths=[1.3, 1.6, 1.9, 1.6],
    )
    doc.add_page_break()

    # ── PART D — already covered ─────────────────────────────────────────
    add_h(doc, "חלק ד׳ — מה כבר התקבל ממך (לא חסר)", level=1,
          color_hex=COCOA, size_pt=20, bottom_border=True)
    add_para(doc, "הבלוקים הבאים כבר נמסרו בשני הקבצים ואינם חסרים:",
             color_hex=INK_DEEP, size_pt=11, space_after=6)
    add_bullets(doc, [
        "אודות גאמוס — הטקסט המלא (תפאורה, קומפלקס, כל אירוע, טכנולוגיה, אנשים).",
        "נרטיב הקולינריה — פילוסופיית המטבח (הכותרת והפתיח של הסקציה).",
        "פסקת הכשרות — ההתאמה לציבור הדתי והחרדי.",
        "טופס יצירת קשר (Lead Form) — כותרת, טקסט, שדות וכפתור (כולל \"שם החברה\").",
        "טאגליין הירו — \"יש אירועים. ויש את גאמוס.\" + כותרת המשנה.",
        "אירועים עסקיים — הנרטיב המלא.",
    ], color_hex=INK_DEEP, size_pt=11)
    add_para(doc, "", space_after=4)
    add_callout(
        doc, "להחלטה נפרדת:",
        "4 בלוקים שמסרת (\"המפרט הטכני\", \"הגן והמדבר / OPEN KITCHEN\", "
        "\"ה-OASIS / בריכה\", \"נוחות ומיקום\") הם תוכן חדש שאין לו עדיין "
        "סקציה באתר — נחליט יחד היכן לשבץ אותם בשלב הבא.",
        "F3E7D9", BRASS_DEEP)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"[docx] saved: {OUT_DOCX}")


if __name__ == "__main__":
    build()
